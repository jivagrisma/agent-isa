"""
Módulo para procesamiento de diferentes tipos de documentos.
"""

import os
import logging
import base64
import hashlib
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple

# Configurar logging
logger = logging.getLogger(__name__)

# Importar bibliotecas para procesamiento de documentos
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    logger.warning("PyPDF2 no está instalado. El soporte para PDF estará limitado.")
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    logger.warning("python-docx no está instalado. El soporte para DOCX estará limitado.")
    DOCX_SUPPORT = False

try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    logger.warning("openpyxl no está instalado. El soporte para XLSX estará limitado.")
    XLSX_SUPPORT = False

class DocumentProcessor:
    """
    Procesador de documentos para diferentes formatos.
    """
    
    def __init__(self):
        """
        Inicializa el procesador de documentos.
        """
        # Crear directorio de caché si no existe
        self.cache_dir = Path(__file__).parent / "cache"
        self.cache_dir.mkdir(exist_ok=True)
        
        logger.info("Procesador de documentos inicializado")
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Procesa un documento y extrae su contenido.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            Diccionario con información del documento
        """
        # Verificar si el archivo existe
        if not os.path.exists(file_path):
            return {"error": f"No se encontró el archivo: {file_path}"}
        
        # Obtener extensión del archivo
        file_ext = Path(file_path).suffix.lower()
        
        # Procesar según el tipo de archivo
        if file_ext == '.pdf':
            return self.process_pdf(file_path)
        elif file_ext == '.docx':
            return self.process_docx(file_path)
        elif file_ext == '.xlsx':
            return self.process_xlsx(file_path)
        elif file_ext in ['.txt', '.md']:
            return self.process_text(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            return self.process_image(file_path)
        else:
            return {"error": f"Tipo de archivo no soportado: {file_ext}"}
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Procesa un archivo PDF.
        
        Args:
            file_path: Ruta al archivo PDF
            
        Returns:
            Diccionario con información del PDF
        """
        if not PDF_SUPPORT:
            return {"error": "La extracción de texto de PDF no está disponible. Instale PyPDF2."}
        
        try:
            # Calcular hash del archivo para caché
            file_hash = self._calculate_file_hash(file_path)
            cache_path = self.cache_dir / f"pdf_{file_hash}.txt"
            
            # Verificar si existe en caché
            if cache_path.exists():
                with open(cache_path, "r", encoding="utf-8") as f:
                    text = f.read()
                logger.info(f"Contenido de PDF recuperado de caché: {file_path}")
                return {
                    "content_type": "application/pdf",
                    "text_content": text,
                    "metadata": {},
                    "from_cache": True
                }
            
            # Extraer texto del PDF
            text = ""
            metadata = {}
            
            with open(file_path, "rb") as file:
                # Crear lector de PDF
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Extraer metadatos
                if pdf_reader.metadata:
                    for key, value in pdf_reader.metadata.items():
                        if value and str(value).strip():
                            # Limpiar nombre de clave (quitar /)
                            clean_key = key.replace('/', '') if isinstance(key, str) else key
                            metadata[clean_key] = str(value)
                
                # Información básica del documento
                info = f"Documento PDF con {num_pages} páginas.\n\n"
                
                # Extraer texto de cada página
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    # Solo añadir separador de página si hay contenido
                    if page_text and page_text.strip():
                        # Añadir número de página para mejor contexto
                        text += f"--- PÁGINA {page_num + 1} ---\n{page_text}\n\n"
            
            # Verificar si se extrajo algún texto
            if not text.strip():
                text = "El PDF no contiene texto extraíble. Puede ser un PDF escaneado o con protección."
            else:
                # Combinar información y texto
                text = info + text
                
                # Guardar en caché
                with open(cache_path, "w", encoding="utf-8") as f:
                    f.write(text)
            
            return {
                "content_type": "application/pdf",
                "text_content": text,
                "metadata": metadata,
                "from_cache": False
            }
                
        except Exception as e:
            logger.error(f"Error al procesar PDF: {e}")
            return {"error": str(e)}
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Procesa un archivo DOCX.
        
        Args:
            file_path: Ruta al archivo DOCX
            
        Returns:
            Diccionario con información del DOCX
        """
        if not DOCX_SUPPORT:
            return {"error": "La extracción de texto de DOCX no está disponible. Instale python-docx."}
        
        try:
            # Calcular hash del archivo para caché
            file_hash = self._calculate_file_hash(file_path)
            cache_path = self.cache_dir / f"docx_{file_hash}.txt"
            
            # Verificar si existe en caché
            if cache_path.exists():
                with open(cache_path, "r", encoding="utf-8") as f:
                    text = f.read()
                logger.info(f"Contenido de DOCX recuperado de caché: {file_path}")
                return {
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "text_content": text,
                    "metadata": {},
                    "from_cache": True
                }
            
            # Extraer texto del DOCX
            doc = docx.Document(file_path)
            
            # Extraer metadatos
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
            }
            
            # Extraer texto de párrafos
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extraer texto de tablas
            tables = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(" | ".join(row_text))
                tables.append("\n".join(table_text))
            
            # Combinar todo el texto
            text = "\n\n".join(paragraphs)
            
            if tables:
                text += "\n\n--- TABLAS ---\n\n" + "\n\n".join(tables)
            
            # Guardar en caché
            with open(cache_path, "w", encoding="utf-8") as f:
                f.write(text)
            
            return {
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "text_content": text,
                "metadata": metadata,
                "from_cache": False
            }
                
        except Exception as e:
            logger.error(f"Error al procesar DOCX: {e}")
            return {"error": str(e)}
    
    def process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """
        Procesa un archivo XLSX.
        
        Args:
            file_path: Ruta al archivo XLSX
            
        Returns:
            Diccionario con información del XLSX
        """
        if not XLSX_SUPPORT:
            return {"error": "La extracción de texto de XLSX no está disponible. Instale openpyxl."}
        
        try:
            # Calcular hash del archivo para caché
            file_hash = self._calculate_file_hash(file_path)
            cache_path = self.cache_dir / f"xlsx_{file_hash}.txt"
            
            # Verificar si existe en caché
            if cache_path.exists():
                with open(cache_path, "r", encoding="utf-8") as f:
                    text = f.read()
                logger.info(f"Contenido de XLSX recuperado de caché: {file_path}")
                return {
                    "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "text_content": text,
                    "metadata": {},
                    "from_cache": True
                }
            
            # Extraer texto del XLSX
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            # Extraer metadatos
            metadata = {
                "sheet_names": workbook.sheetnames
            }
            
            # Extraer texto de cada hoja
            sheets_text = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Obtener dimensiones de la hoja
                min_row, min_col = 1, 1
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                # Extraer encabezados
                headers = []
                for col in range(min_col, max_col + 1):
                    cell_value = sheet.cell(row=min_row, column=col).value
                    headers.append(str(cell_value) if cell_value is not None else "")
                
                # Extraer datos
                rows = []
                for row in range(min_row + 1, max_row + 1):
                    row_data = []
                    for col in range(min_col, max_col + 1):
                        cell_value = sheet.cell(row=row, column=col).value
                        row_data.append(str(cell_value) if cell_value is not None else "")
                    rows.append(" | ".join(row_data))
                
                # Combinar encabezados y datos
                sheet_text = f"--- HOJA: {sheet_name} ---\n"
                sheet_text += " | ".join(headers) + "\n"
                sheet_text += "\n".join(rows)
                
                sheets_text.append(sheet_text)
            
            # Combinar texto de todas las hojas
            text = "\n\n".join(sheets_text)
            
            # Guardar en caché
            with open(cache_path, "w", encoding="utf-8") as f:
                f.write(text)
            
            return {
                "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "text_content": text,
                "metadata": metadata,
                "from_cache": False
            }
                
        except Exception as e:
            logger.error(f"Error al procesar XLSX: {e}")
            return {"error": str(e)}
    
    def process_text(self, file_path: str) -> Dict[str, Any]:
        """
        Procesa un archivo de texto.
        
        Args:
            file_path: Ruta al archivo de texto
            
        Returns:
            Diccionario con información del archivo de texto
        """
        try:
            # Leer archivo de texto
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            
            return {
                "content_type": "text/plain",
                "text_content": text,
                "metadata": {},
                "from_cache": False
            }
                
        except Exception as e:
            logger.error(f"Error al procesar archivo de texto: {e}")
            return {"error": str(e)}
    
    def process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Procesa una imagen.
        
        Args:
            file_path: Ruta a la imagen
            
        Returns:
            Diccionario con información de la imagen
        """
        try:
            # Obtener tipo MIME
            mime_type = self._get_mime_type(file_path)
            
            # Leer imagen como base64
            with open(file_path, "rb") as f:
                image_data = f.read()
                base64_image = base64.b64encode(image_data).decode("utf-8")
            
            return {
                "content_type": mime_type,
                "image_data": base64_image,
                "metadata": {},
                "from_cache": False
            }
                
        except Exception as e:
            logger.error(f"Error al procesar imagen: {e}")
            return {"error": str(e)}
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """
        Calcula el hash MD5 de un archivo.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            Hash MD5 del archivo
        """
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _get_mime_type(self, file_path: str) -> str:
        """
        Determina el tipo MIME de un archivo basado en su extensión.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            Tipo MIME
        """
        extension = Path(file_path).suffix.lower()
        
        mime_types = {
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".webp": "image/webp"
        }
        
        return mime_types.get(extension, "application/octet-stream")
